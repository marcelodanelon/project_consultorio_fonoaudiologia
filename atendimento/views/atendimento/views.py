from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.forms import inlineformset_factory
from django.shortcuts import get_object_or_404
from django.urls import reverse
from django.http import JsonResponse
from django.contrib import messages
from datetime import date
from django.db.models import Q, Case, When, Value, CharField
from atendimento.forms import AtendimentoForm, RegulagemForm, ContatosTelefonicosForm
from atendimento.models import AtendimentoModel, RegulagemModel, ContatosTelefonicosModel, AudiometriaModel
from home.models import ClientModel, ProfessionalModel, LocalModel
from estoque.models import MovimentacaoInsumoModel, ItensMovimentacaoInsumoModel
from agendamento.models import AgendamentoModel
from docx import Document
from io import BytesIO
from django.http import HttpResponse
import os
import pythoncom
import win32com.client
from tempfile import NamedTemporaryFile
from docx.shared import Inches

def generate_word_document(data, url, idRegistro=None):
    doc = Document(url)

    # Função para substituir campos de texto em um parágrafo
    def replace_text_in_paragraph(paragraph, data):
        for run in paragraph.runs:
            for key, valor in data.items():
                if valor is None:
                    run.text = run.text.replace(f'[{key}]', ' ')
                else:
                    if valor is True:
                        run.text = run.text.replace(f'[{key}]', 'X')
                    elif valor is False:
                        run.text = run.text.replace(f'[{key}]', ' ')
                    else:
                        run.text = run.text.replace(f'[{key}]', str(data[key]))

    def replace_image_in_paragraph(paragraph, image_path):
        if os.path.exists(image_path):
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(2.0))
        else:
            print(f"Arquivo de imagem '{image_path}' não encontrado.")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data)
        
    for paragraph in doc.paragraphs:
        if idRegistro:
            if 'OE' in paragraph.text: 
                replace_image_in_paragraph(paragraph, f'utils/data_files/temp_images_audiometria/{idRegistro}_plano_cartesiano_OE.png') 
            elif 'OD' in paragraph.text: 
                replace_image_in_paragraph(paragraph, f'utils/data_files/temp_images_audiometria/{idRegistro}_plano_cartesiano_OD.png')  
            else:
                replace_text_in_paragraph(paragraph, data)

    if 'ATENDIMENTO' in url:
        for reg in data['anamneses_objs']:
            doc.add_paragraph('REGULAGENS')
            doc.add_paragraph('-' * 100)
            
            data_paragraph = doc.add_paragraph()
            data_paragraph.add_run('Data: ').bold = True
            data_paragraph.add_run(f'{str(data["aDataAte"])}')

            od_paragraph = doc.add_paragraph()
            od_paragraph.add_run('OD:').bold = True
            od_paragraph.add_run(f' {reg.aAjustOD}')

            oe_paragraph = doc.add_paragraph()
            oe_paragraph.add_run('OE:').bold = True
            oe_paragraph.add_run(f' {reg.aAjustOE}')

            obs_paragraph = doc.add_paragraph()
            obs_paragraph.add_run('Obs:').bold = True
            obs_paragraph.add_run(f' {reg.aAObserv}')

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def save_as_pdf(docx_file):
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(docx_file)
        pdf_file = docx_file.replace(".docx", ".pdf")
        doc.SaveAs(pdf_file, FileFormat=17)
        doc.Close()
        word.Quit()
        return pdf_file
    except Exception as e:
        print(f"Error converting to PDF: {e}")
    finally:
        pythoncom.CoUninitialize()

@login_required(login_url='home:loginUser')
def getJSONclient(request):
    if request.GET.get('searchClient'):
        q = request.GET.get('searchClient')
        model = list(ClientModel.objects.filter(Q(first_name__contains=q) | Q(last_name__contains=q)).values('id', 'first_name','last_name', 'born', 'document1'))
        return JsonResponse(data={'results': model})

@login_required(login_url='home:loginUser')
def index(request):
    # Criação primeiro gráfico por Atendimento
    count = int(ProfessionalModel.objects.all().count())
    professionals = list(ProfessionalModel.objects.all())
    data_points1 = []
    for i in range(count):
        professional = AtendimentoModel.objects.filter(aProfessional=professionals[i])
        if not professional:
            data_points1.append({'label': 'Sem registros', "y": 0})
            break
        else:
            data_points1.append({'label': str(professional.first().aProfessional), "y": professional.count()})

    # Criação segundo gráfico por Unidade
    count = int(LocalModel.objects.all().count())
    locais = list(LocalModel.objects.all())
    data_points2 = []
    for i in range(count):
        local = AtendimentoModel.objects.filter(aLocal=locais[i])
        if not local:
            data_points2.append({'label': 'Sem registros', "y": 0})
            break
        else:
            data_points2.append({'label': locais[i].name, "y": local.count()})

    # Criação terceiro gráfico por mês no ano atual (Anamnese)
    data_meses = ['Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho', 'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro']
    data_points3 = []
    for i in range(12):
        meses = AtendimentoModel.objects.filter(aDataAte__month=i).filter(aDataAte__year=date.today().year)
        if meses.count() != 0:
            data_points3.append({'label': data_meses[i], "y": meses.count()})

    # Criação quarto gráfico por mês no ano atual (1º Atendimento)
    data_points4 = []
    for i in range(12):
        meses = AtendimentoModel.objects.filter(aDataPri__month=i).filter(aDataPri__year=date.today().year)
        if meses.count() != 0:
            data_points4.append({'label': data_meses[i], "y": meses.count()})

    context = {
        'title': 'Home',
        'name_module': 'Home',
        'data_points' : data_points1,
        'data_points2' : data_points2,
        'data_points3' : data_points3,
        'data_points4' : data_points4,
        'name_module': 'Atendimento',
        'title': 'Atendimento',
    }

    return render(
        request,
        'atendimento/atendimento/index.html',
        context
    )

@login_required(login_url='home:loginUser')
def historicoAtendimento(request):
    form_atendimento = AtendimentoForm()

    try:
        search = int(request.GET.get('searchClient'))
        client = ClientModel.objects.filter(pk=search).get()               
    except:
        search = None
        client = None

    client_atendimentos = AtendimentoModel.objects.filter(aClient=client)
    client_anamneses = RegulagemModel.objects.none()
    for i in range(client_atendimentos.count()):
        client_anamneses = client_anamneses.union(RegulagemModel.objects.filter(aIDAtend=client_atendimentos[i]))
    client_telefonemas = ContatosTelefonicosModel.objects.filter(aIDAtend__aClient=client)

    from django.db.models import F

    client_saida = MovimentacaoInsumoModel.objects.exclude(eClient=None).filter(eClient=client)


    client_saida_com_nome_insumo = []

    for movimentacao in client_saida:
        itens_insumo = ItensMovimentacaoInsumoModel.objects.filter(movimentacao=movimentacao)
        
        for item in itens_insumo:
            descricao_insumo = item.insumo.descricao
            client_saida_com_nome_insumo.append({
                'movimentacao_id': movimentacao.id,
                'nome_insumo': descricao_insumo
            })
    for client_saida_item in client_saida_com_nome_insumo:
        print(f'Movimentação ID: {client_saida_item["movimentacao_id"]}, Nome do Insumo: {client_saida_item["nome_insumo"]}')

    client_saida = MovimentacaoInsumoModel.objects.exclude(eClient=None).filter(eClient=client)

    client_saida = client_saida.annotate(
        nome_insumo=Case(
            When(itensmovimentacaoinsumomodel__movimentacao=F('id'), then=F('itensmovimentacaoinsumomodel__insumo__descricao')),
            default=Value(''), output_field=CharField()
        )
    )

    client_audiometria = AudiometriaModel.objects.filter(aClient=client)

    client_agendamentos = AgendamentoModel.objects.filter(aClient=client)
    for i in client_agendamentos:
        i.agDataAg = i.agDataAg.strftime("%d/%m/%Y")

    context = {
        'form_atendimento': form_atendimento,
        'atendimentos': client_atendimentos,
        'telefonemas': client_telefonemas,
        'anamneses': client_anamneses,
        'saidasEstoque': client_saida,
        'audiometrias': client_audiometria,
        'agendamentos': client_agendamentos,
        'name_module': 'Atendimento',
        'title': 'Atendimento',
    }

    return render(
        request,
        'atendimento/atendimento/historicoAtendimento.html',
        context
    )

@login_required(login_url='home:loginUser')
def atendimento(request):
    form_action = reverse('atendimento:atendimento')
    try:
        search = int(request.GET.get('searchClient'))
        client = ClientModel.objects.filter(pk=search)
    except:
        search = None
        client = None
    atendimentoForm = AtendimentoForm()
    telefonemasFormset = inlineformset_factory(AtendimentoModel, ContatosTelefonicosModel, form=ContatosTelefonicosForm, extra=0, can_delete=True, min_num=1)
    telefonemasForm = telefonemasFormset(instance=AtendimentoModel())
    regulagensFormset = inlineformset_factory(AtendimentoModel, RegulagemModel, form=RegulagemForm, extra=0, can_delete=True, min_num=1)
    regulagemForm = regulagensFormset(instance=AtendimentoModel())

    if request.method == 'POST':
        agendamento_id = request.POST.get('agendamentoId')
        if agendamento_id:
            agendamento_id = int(agendamento_id)
            agendamento = get_object_or_404(AgendamentoModel, id=agendamento_id)
            agendamento.agSituac = 'atendido'
            agendamento.save()
        
        atendimentoForm = AtendimentoForm(request.POST, request.FILES, instance=AtendimentoModel())
        regulagemForm = regulagensFormset(request.POST, request.FILES, instance=AtendimentoModel())
        telefonemasForm = telefonemasFormset(request.POST, request.FILES, instance=AtendimentoModel())

        if atendimentoForm.is_valid() and regulagemForm.is_valid() and telefonemasForm.is_valid():            
            formAten_instance = atendimentoForm.save()      
            formTel_instances = telefonemasForm.save(commit=False) 
            formReg_instances = regulagemForm.save(commit=False) 
            for item in formTel_instances:
                item.aIDAtend = formAten_instance
                item.save()
            for item in formReg_instances:
                item.aIDAtend = formAten_instance
                item.save()
            messages.success(request, 'Atendimento gravado com sucesso!')

            response = HttpResponse(status=204)
            response['id_registro'] = str(formAten_instance.id)
            return response

    context={
        'atendimentoForm': atendimentoForm,
        'telefonemasForm': telefonemasForm,
        'regulagemForm': regulagemForm,
        'form_action': form_action,
        'client': client,
        'name_module': 'Atendimento',
        'title': 'Atendimento',
    }

    return render(
        request,
        'atendimento/atendimento/atendimento.html',
        context
    )

def download_documento_atendimento(request):
    id_registro = request.GET.get('registro')
    documento = request.GET.get('documento')
    atendimento_obj = get_object_or_404(AtendimentoModel, id=id_registro)
    anamneses_objs = RegulagemModel.objects.filter(aIDAtend=id_registro)
    
    url = ''
    data = {
        'aProfessional': atendimento_obj.aProfessional,
        'aDataAte': atendimento_obj.aDataAte,
        'aLocal': atendimento_obj.aLocal,
        'aClient': atendimento_obj.aClient,
        'aClient.age': atendimento_obj.aClient.age,
        'aClient.street': atendimento_obj.aClient.street,
        'aClient.district': atendimento_obj.aClient.district,
        'aClient.born': atendimento_obj.aClient.born,
        'aClient.city': atendimento_obj.aClient.city,
        'aClient.zipcode': atendimento_obj.aClient.zipcode,
        'aClient.phone1': atendimento_obj.aClient.phone1,
        'aClient.phone2': atendimento_obj.aClient.phone2,
        'aClient.profession': atendimento_obj.aClient.profession,
        'aClient.document1': atendimento_obj.aClient.document1,
        'aClient.document2': atendimento_obj.aClient.document2,
        'aConhece': atendimento_obj.aConhece,
        'aDifiEsc': atendimento_obj.aDifiEsc,
        'aDifiPio': atendimento_obj.aDifiPio,
        'aOuviMel': atendimento_obj.aOuviMel,
        'aPessFam': atendimento_obj.aPessFam,
        'aTrabRui': atendimento_obj.aTrabRui,
        'aTelevis': atendimento_obj.aTelevis,
        'aTeleFix': atendimento_obj.aTeleFix,
        'aTeleCel': atendimento_obj.aTeleCel,
        'aConvGru': atendimento_obj.aConvGru,
        'aConvRui': atendimento_obj.aConvRui,
        'aFalaBai': atendimento_obj.aFalaBai,
        'aFaladis': atendimento_obj.aFaladis,
        'aCineTea': atendimento_obj.aCineTea,
        'aPaleSal': atendimento_obj.aPaleSal,
        'aOutrDif': atendimento_obj.aOutrDif,
        'aZumbido': atendimento_obj.aZumbido,
        'aCoceira': atendimento_obj.aCoceira,
        'aOtiteOO': atendimento_obj.aOtiteOO,
        'aDorOOOO': atendimento_obj.aDorOOOO,
        'aCiruOuv': atendimento_obj.aCiruOuv,
        'aTimpPer': atendimento_obj.aTimpPer,
        'aSensTam': atendimento_obj.aSensTam,
        'aOutrOuv': atendimento_obj.aOutrOuv,
        'aJaTesAp': atendimento_obj.aJaTesAp,
        'aQualApa': atendimento_obj.aQualApa,
        'aApaIndi': atendimento_obj.aApaIndi,
        'aValApar': atendimento_obj.aValApar,
        'aLadoInd': atendimento_obj.aLadoInd,
        'aFormPag': atendimento_obj.aFormPag,
        'aSaiTest': atendimento_obj.aSaiTest,
        'aRetTest': atendimento_obj.aRetTest,
        'aComClik': atendimento_obj.aComClik,
        'aSemClik': atendimento_obj.aSemClik,
        'aClikOOD': atendimento_obj.aClikOOD,
        'aClikOOE': atendimento_obj.aClikOOE,
        'aTuboOOD': atendimento_obj.aTuboOOD,
        'aTuboOOE': atendimento_obj.aTuboOOE,
        'aReceOOD': atendimento_obj.aReceOOD,
        'aReceOOE': atendimento_obj.aReceOOE,
        'aJaUsoAp': atendimento_obj.aJaUsoAp,
        'aMarcaOO': atendimento_obj.aMarcaOO,
        'aTempoOO': atendimento_obj.aTempoOO,
        'aDataAte': atendimento_obj.aDataAte,
        'anamneses_objs': anamneses_objs,
    }

    match documento:
        case 'fichaAtendimento':
            doc = 'FICHA DE ATENDIMENTO.docx'
            url = 'utils\\docs\\FICHA DE ATENDIMENTO.docx'
        case 'declaracaoComparecimento':
            doc = 'DECLARACAO DE COMPARECIMENTO.docx'
            url = 'utils\\docs\\DECLARACAO DE COMPARECIMENTO.docx'

    word_document = generate_word_document(data, url)

    temp_file = NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file.write(word_document.getvalue())
    temp_file.close()

    pdf_file_path = save_as_pdf(temp_file.name)

    with open(pdf_file_path, "rb") as pdf_file:
        response = HttpResponse(pdf_file.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename={doc.replace(".docx", ".pdf")}'

    os.remove(temp_file.name)

    return response

